
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from docx import Document
from pathlib import Path

# Configurações do Chrome (modo headless opcional)
chrome_options = Options()
chrome_options.add_argument("--start-maximized")
# chrome_options.add_argument("--headless")  # Ative se quiser sem abrir janela

driver = webdriver.Chrome(options=chrome_options)

TIME = 2

try:
    # Passo 1: Acessa a página de login
    url_login = "https://www.valadares.mg.gov.br:8002/login.aspx?ReturnUrl=%2fdefault.aspx"
    driver.get(url_login)
    time.sleep(TIME)  # Aguarda carregamento da página

    # Passo 2: Preenche o formulário de login
    driver.find_element(By.NAME, "txt_usuario").send_keys("jose.coelho")
    driver.find_element(By.NAME, "txt_senha").send_keys("Jose@2025")
    driver.find_element(By.NAME, "btn_entrar").click()

    # Passo 3: Aguarda carregamento da página principal
    WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.ID, "centraliza"))
    )
    print("Login realizado com sucesso!")

    # # Passo 4: Tenta abrir o menu lateral, se necessário
    # try:
    #     toggle_menu = driver.find_element(By.ID, "m_aside_left_offcanvas_toggle")
    #     toggle_menu.click()
    #     time.sleep(1)
    # except:
    #     pass  # ignora se não encontrar o botão
    #
    # # Passo 5: Clica no botão "Ouvidoria"
    # ouvidoria_button = WebDriverWait(driver, 10).until(
    #     EC.element_to_be_clickable((By.XPATH, "/html/body/div[1]/div[2]/div[1]/div/ul/li[7]/a/span"))
    # )
    # ouvidoria_button.click()
    # print("Botão 'Ouvidoria' clicado com sucesso!")
    #
    # # Passo 6: Clica no botão "Manifestação"
    # manifestacao_button = WebDriverWait(driver, 10).until(
    #     EC.element_to_be_clickable((By.XPATH, "/html/body/div[1]/div[2]/div[1]/div/ul/li[7]/div/ul/li[4]/a/span"))
    # )
    # manifestacao_button.click()
    # print("Botão 'Manifestação' clicado com sucesso!")
    while True:
        url = 'https://www.valadares.mg.gov.br:8002/ouv_ger.aspx?formType=1'
        driver.get(url)
        time.sleep(TIME)

        print('entrando no site!')
        # Passo 7: Aguarda e preenche o campo de busca
        search_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "input[aria-controls='gvw_ouvidoria']"))
        )
        manifestacao_cod = input('MANIFESTAÇÃO: ')
        if manifestacao_cod.lower() == 'sair':  
            break
        else:
            pass
        search_input.send_keys(manifestacao_cod)  # <- Altere esse texto como quiser
        print("Texto inserido no campo de busca!")

        # Opcional: Pressiona Enter
        search_input.send_keys(Keys.RETURN)
        print("Busca acionada!")
        time.sleep(5)
        # Passo 8: Clica no botão "Visualizar" (ícone de pasta)
        visualizar_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,"/html/body/div[1]/div/div[2]/div[2]/div/div/div/div["
                                                 "2]/div/div/div/div/div[2]/div/table/tbody/tr/td[11]/a"))
        )
        visualizar_button.click()
        print("Botão 'Visualizar' clicado com sucesso!")

        # Passo 9: Extrai os dados da manifestação
        time.sleep(3)  # espera o carregamento da página

        identidade = driver.find_element(By.ID, "lbl_IdentidadeManifestante").text.strip()

        if identidade.lower() == "sigilosa":
            nome = 'Sigiloso'
            data = driver.find_element(By.ID, "lbl_dtManifesto").text.strip()
            numero = "Sigiloso"
            origem = driver.find_element(By.ID, "lbl_dmOrigem").text.strip()
            tipo = driver.find_element(By.ID, "lbl_dsTipoManifesto").text.strip()
            mensagem = driver.find_element(By.ID, "lbl_dsManifesto").text.strip()

        else:
            nome = driver.find_element(By.ID, "lbl_nmManifestante").text.strip()
            data = driver.find_element(By.ID, "lbl_dtManifesto").text.strip()
            numero = driver.find_element(By.ID, "lbl_cdControle").text.strip()
            origem = driver.find_element(By.ID, "lbl_dmOrigem").text.strip()
            tipo = driver.find_element(By.ID, "lbl_dsTipoManifesto").text.strip()
            mensagem = driver.find_element(By.ID, "lbl_dsManifesto").text.strip()

        # Clica no botão "Mais informações"
        mais_info_btn = driver.find_element(By.XPATH, '/html/body/div[1]/div/div[2]/form/div[3]/div/div/div/div/div[1]/div[3]/div[2]/div/div/div/div[2]/div/table/tbody/tr[4]/td[6]/a')
        driver.execute_script("arguments[0].click();", mais_info_btn)
        print('Mais informações da resposta conclusiva encontradas')
        time.sleep(2)
        #         from selenium.webdriver.common.by import By
        #         from selenium.webdriver.support.ui import WebDriverWait
        #         from selenium.webdriver.support import expected_conditions as EC
        #
        #         rows = driver.find_elements(By.CSS_SELECTOR, "#gvw_historico tbody tr")
        #
        #         for index, row in enumerate(rows):
        #             cells = row.find_elements(By.TAG_NAME, "td")
        #             for cell in cells:
        #                 if cell.text.strip() == "Resposta conclusiva do supervisor":
        #                     try:
        #                         # Aguarda até que o botão esteja presente e clicável
        #                         link = WebDriverWait(row, 10).until(
        #                             EC.element_to_be_clickable((By.CSS_SELECTOR, "td.no-clickable a"))
        #                         )
        #                         link.click()
        #                         print("[SUCESSO] Botão correto clicado com sucesso!")
        #                         break
        #                     except Exception as e:
        #                         print(f"[ERRO] Não foi possível clicar no botão: {e}")
        #             else:
        #                 continue
        #             break

        # Extrai o texto adicional da resposta conclusiva
        # resposta_conclusiva = driver.find_element(By.XPATH, '/html/body/div[1]/div/div[2]/form/div[3]/div/div/div/div/div[1]/div[3]/div[3]/div/div/div/div[2]/div/p[5]').text.strip()
        resposta_conclusiva = driver.find_element(By.CSS_SELECTOR, "#p_dsResposta").text.strip()
        # data_conclusiva = driver.find_element(By.XPATH,'/html/body/div[1]/div/div[2]/form/div[3]/div/div/div/div/div[1]/div[3]/div[3]/div/div/div/div[2]/div/p[1]').text.strip()
        data_conclusiva = driver.find_element(By.CSS_SELECTOR, "#p_dtFluxo").text.strip()

        # Impressão com a nova ordem solicitada
        dados = {"Data": f"{data}",
        "Nome": f"{nome}",
        "Número da ocorrência": f"{numero}",
        "Origem": f"{origem}",
        "Tipo": f"{tipo}",
        "Mensagem": f"{mensagem}",
        f"Resposta conclusiva em {data_conclusiva.split(' ')[1]}": f"{resposta_conclusiva[9:]}\n"
                 }
        print('escrevendo arquivo')
        arquivo = "resultado.docx"
        caminho = Path(arquivo)

        try:
            # Se o documento já existir, ele é carregado
            if caminho.is_file():
                documento = Document(arquivo)
            else:
                documento = Document()
                print('Gravando dados')
        except Exception as e:
            print("Erro ao abrir ou criar o documento:", e)
            # Se houver erro, cria um documento novo como fallback
            documento = Document()

        # Para cada par chave-valor, adiciona um parágrafo com a chave em negrito
        for chave, valor in dados.items():
            paragrafo = documento.add_paragraph()
            run_chave = paragrafo.add_run(f"{chave}: ")
            run_chave.bold = True
            paragrafo.add_run(valor)

        # Salva ou atualiza o documento
        documento.save(arquivo)
        print('Dados gravados')
    # except Exception as e:
    #     print("Erro durante o processo:", e)

finally:
    driver.quit()
